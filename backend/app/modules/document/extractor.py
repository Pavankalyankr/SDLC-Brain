"""
SDLC Brain — Document Text Extractor

Extracts plain text from uploaded SOW documents.
Supports: PDF (via PyMuPDF), DOCX (via python-docx), plain text.
"""

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


async def extract_text(file_path: str, content_type: str) -> tuple[str, int | None]:
    """
    Extract text content from a document file.

    Returns:
        (extracted_text, page_count)
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    if content_type == "application/pdf" or path.suffix.lower() == ".pdf":
        return _extract_pdf(path)
    elif content_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ) or path.suffix.lower() in (".docx", ".doc"):
        return _extract_docx(path)
    elif content_type.startswith("text/") or path.suffix.lower() in (".txt", ".md", ".rst"):
        return _extract_text(path)
    else:
        logger.warning(f"Unsupported content type: {content_type}")
        return "", None


def _extract_pdf(path: Path) -> tuple[str, int | None]:
    """Extract text from PDF using PyMuPDF."""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(str(path))
        pages = []
        for page in doc:
            pages.append(page.get_text())
        doc.close()

        text = "\n\n".join(pages)
        return text.strip(), len(pages)
    except ImportError:
        logger.warning("PyMuPDF not installed. Install with: pip install pymupdf")
        return _extract_text(path)
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return "", None


def _extract_docx(path: Path) -> tuple[str, int | None]:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document

        doc = Document(str(path))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        text = "\n\n".join(paragraphs)
        return text.strip(), None
    except ImportError:
        logger.warning("python-docx not installed. Install with: pip install python-docx")
        return _extract_text(path)
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return "", None


def _extract_text(path: Path) -> tuple[str, int | None]:
    """Extract text from plain text files."""
    try:
        text = path.read_text(encoding="utf-8")
        return text.strip(), None
    except UnicodeDecodeError:
        text = path.read_text(encoding="latin-1")
        return text.strip(), None
    except Exception as e:
        logger.error(f"Text extraction failed: {e}")
        return "", None
