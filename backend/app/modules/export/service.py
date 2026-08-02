"""
SDLC Brain — Export Service

Handles generating PDF, DOCX, and Markdown exports for project artifacts.
"""

import io
import json
from datetime import datetime
from typing import Any

from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

class ExportService:
    def export_agile_to_pdf(self, module_type: str, items: list[Any], project_name: str = "Project") -> io.BytesIO:
        """
        Generates a professional PDF document from a list of agile artifacts.
        Returns an in-memory BytesIO buffer containing the PDF data.
        """
        buffer = io.BytesIO()
        
        # Use landscape A4 to give tables more width
        doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
        elements = []
        
        styles = getSampleStyleSheet()
        title_style = styles["Title"]
        normal_style = styles["Normal"]
        
        # Header
        title = f"{project_name} - {module_type.capitalize()} Export"
        elements.append(Paragraph(title, title_style))
        elements.append(Paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", normal_style))
        elements.append(Spacer(1, 20))
        
        if not items:
            elements.append(Paragraph("No items found.", normal_style))
            doc.build(elements)
            buffer.seek(0)
            return buffer
            
        # Determine table columns based on module type
        if module_type == "requirements":
            headers = ["ID", "Title", "Category", "Priority", "Status", "Description"]
        elif module_type == "epics":
            headers = ["ID", "Title", "Status", "Description"]
        elif module_type == "features":
            headers = ["ID", "Title", "Status", "Description"]
        elif module_type == "stories":
            headers = ["ID", "Title", "Points", "Sprint", "Status", "Acceptance Criteria"]
        else:
            headers = ["ID", "Title", "Status", "Description"]

        data = [headers]
        
        for item in items:
            # Safely get attributes since they might be SQLAlchemy models or Pydantic schemas
            def get_val(obj, key, default=""):
                val = getattr(obj, key, default)
                return str(val) if val is not None else default
            
            # Format text in paragraphs to allow word wrapping in table cells
            title_para = Paragraph(get_val(item, "title"), normal_style)
            status_para = Paragraph(get_val(item, "status").upper(), normal_style)
            short_id = get_val(item, "id")[:8]
            
            if module_type == "requirements":
                desc_para = Paragraph(get_val(item, "description"), normal_style)
                data.append([short_id, title_para, get_val(item, "category"), get_val(item, "priority"), status_para, desc_para])
            elif module_type == "epics" or module_type == "features":
                desc_para = Paragraph(get_val(item, "description"), normal_style)
                data.append([short_id, title_para, status_para, desc_para])
            elif module_type == "stories":
                criteria_para = Paragraph(get_val(item, "acceptance_criteria", get_val(item, "description")), normal_style)
                data.append([short_id, title_para, get_val(item, "story_points"), get_val(item, "sprint"), status_para, criteria_para])
            else:
                desc_para = Paragraph(get_val(item, "description"), normal_style)
                data.append([short_id, title_para, status_para, desc_para])

        # Define column widths based on module type
        # Landscape A4 is approx 842 points wide, minus 60 margins = 782 available points
        if module_type == "requirements":
            col_widths = [60, 150, 70, 50, 60, 392]
        elif module_type == "epics" or module_type == "features":
            col_widths = [60, 200, 60, 462]
        elif module_type == "stories":
            col_widths = [60, 150, 40, 60, 60, 412]
        else:
            col_widths = None # Auto-distribute

        table = Table(data, colWidths=col_widths, repeatRows=1)
        
        # Add styling to the table
        style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#f1f5f9")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor("#0f172a")),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('TOPPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor("#334155")),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor("#e2e8f0")),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")])
        ])
        
        table.setStyle(style)
        elements.append(table)
        
        doc.build(elements)
        buffer.seek(0)
        return buffer

    def export_agile(self, module_type: str, items: list[Any], format_type: str, project_name: str = "Project") -> tuple[io.BytesIO, str, str]:
        if format_type == "pdf":
            buf = self.export_agile_to_pdf(module_type, items, project_name)
            return buf, "application/pdf", "pdf"
        elif format_type == "json":
            def get_dict(item):
                d = {}
                for key in ["id", "title", "description", "category", "priority", "status", "acceptance_criteria", "story_points", "sprint", "version"]:
                    if hasattr(item, key):
                        d[key] = str(getattr(item, key)) if getattr(item, key) is not None else None
                return d
            data = [get_dict(i) for i in items]
            buffer = io.BytesIO()
            buffer.write(json.dumps({"project": project_name, "module": module_type, "generated_at": datetime.now().isoformat(), "count": len(data), "items": data}, indent=2).encode("utf-8"))
            buffer.seek(0)
            return buffer, "application/json", "json"
        elif format_type == "md":
            lines = [f"# {project_name} — {module_type.capitalize()} Export\n", f"*Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n\n---\n"]
            if not items:
                lines.append("*No items found.*\n")
            for item in items:
                lines.append(f"## {getattr(item, 'title', 'Untitled')} (`{getattr(item, 'status', 'draft').upper()}`)")
                if getattr(item, "priority", None):
                    lines.append(f"**Priority:** {getattr(item, 'priority')} | **Category:** {getattr(item, 'category', 'N/A')}\n")
                if getattr(item, "story_points", None) is not None:
                    lines.append(f"**Points:** {getattr(item, 'story_points')} | **Sprint:** {getattr(item, 'sprint', 'N/A')}\n")
                if getattr(item, "description", None):
                    lines.append(f"{getattr(item, 'description')}\n")
                if getattr(item, "acceptance_criteria", None):
                    lines.append(f"### Acceptance Criteria\n{getattr(item, 'acceptance_criteria')}\n")
                lines.append("---\n")
            buffer = io.BytesIO()
            buffer.write("\n".join(lines).encode("utf-8"))
            buffer.seek(0)
            return buffer, "text/markdown", "md"
        else: # docx
            doc_obj = Document()
            doc_obj.add_heading(f"{project_name} - {module_type.capitalize()} Export", 0)
            doc_obj.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            if not items:
                doc_obj.add_paragraph("No items found.")
            else:
                for item in items:
                    doc_obj.add_heading(str(getattr(item, "title", "Untitled")), level=1)
                    p = doc_obj.add_paragraph()
                    p.add_run(f"Status: {str(getattr(item, 'status', 'draft')).upper()}").bold = True
                    if getattr(item, "description", None):
                        doc_obj.add_paragraph(str(getattr(item, "description")))
                    if getattr(item, "acceptance_criteria", None):
                        doc_obj.add_heading("Acceptance Criteria", level=2)
                        doc_obj.add_paragraph(str(getattr(item, "acceptance_criteria")))
            buffer = io.BytesIO()
            doc_obj.save(buffer)
            buffer.seek(0)
            return buffer, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"

    def export_architecture(self, module_type: str, items: list[Any], format_type: str, project_name: str = "Project") -> tuple[io.BytesIO, str, str]:
        """
        Generates PDF, DOCX, Markdown, or JSON export for architecture items.
        Returns (BytesIO buffer, media_type, file_extension).
        """
        if format_type == "json":
            def get_dict(item):
                d = {}
                for key in ["id", "title", "description", "architecture_type", "components", "tech_stack",
                            "mermaid_diagram", "method", "path", "summary", "request_body", "response_body",
                            "status_codes", "service", "table_name", "columns", "relationships", "indexes",
                            "status", "version", "source_type", "source_id"]:
                    if hasattr(item, key):
                        val = getattr(item, key)
                        if key in ["components", "tech_stack", "request_body", "response_body", "status_codes", "columns", "relationships", "indexes"] and isinstance(val, str):
                            try:
                                val = json.loads(val)
                            except Exception:
                                pass
                        d[key] = val
                return d
            data = [get_dict(item) for item in items]
            buffer = io.BytesIO()
            buffer.write(json.dumps({
                "project": project_name,
                "module": module_type,
                "generated_at": datetime.now().isoformat(),
                "count": len(data),
                "items": data
            }, indent=2).encode("utf-8"))
            buffer.seek(0)
            return buffer, "application/json", "json"

        elif format_type == "md":
            lines = [f"# {project_name} — {module_type.replace('-', ' ').title()} Export\n",
                     f"*Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n\n---\n"]
            if not items:
                lines.append("*No items found.*\n")
            for item in items:
                if module_type in ["designs", "system-design", "system_design"]:
                    lines.append(f"## {getattr(item, 'title', 'Untitled Design')} ({getattr(item, 'architecture_type', 'System')})")
                    lines.append(f"**Status:** {getattr(item, 'status', 'draft').upper()} | **Version:** {getattr(item, 'version', 1)}\n")
                    if getattr(item, "description", ""):
                        lines.append(f"### Description\n{getattr(item, 'description')}\n")
                    tech = getattr(item, "tech_stack", "{}")
                    try:
                        tech_dict = json.loads(tech) if isinstance(tech, str) else tech
                        if tech_dict and isinstance(tech_dict, dict):
                            lines.append("### Recommended Tech Stack")
                            for k, v in tech_dict.items():
                                lines.append(f"- **{k.title()}:** {v}")
                            lines.append("")
                    except Exception:
                        pass
                    comps = getattr(item, "components", "[]")
                    try:
                        comps_list = json.loads(comps) if isinstance(comps, str) else comps
                        if comps_list and isinstance(comps_list, list):
                            lines.append("### Architectural Components")
                            lines.append("| Name | Type | Tech | Description |")
                            lines.append("| :--- | :--- | :--- | :--- |")
                            for c in comps_list:
                                lines.append(f"| {c.get('name', '')} | {c.get('type', '')} | {c.get('tech', '')} | {c.get('description', '')} |")
                            lines.append("")
                    except Exception:
                        pass
                    diag = getattr(item, "mermaid_diagram", None)
                    if diag:
                        lines.append("### Architecture Diagram\n```mermaid\n" + diag.strip() + "\n```\n")
                    lines.append("---\n")

                elif module_type in ["apis", "api-contracts", "api_contracts"]:
                    method = getattr(item, "method", "GET").upper()
                    path = getattr(item, "path", "/")
                    lines.append(f"## `[{method}]` {path}")
                    lines.append(f"**Summary:** {getattr(item, 'summary', '')} | **Status:** {getattr(item, 'status', 'draft').upper()}\n")
                    if getattr(item, "description", ""):
                        lines.append(f"{getattr(item, 'description')}\n")
                    req = getattr(item, "request_body", None)
                    if req and req != "null" and req != "None":
                        lines.append("### Request Body\n```json\n" + str(req) + "\n```\n")
                    res = getattr(item, "response_body", None)
                    if res and res != "null" and res != "None":
                        lines.append("### Response Body\n```json\n" + str(res) + "\n```\n")
                    lines.append("---\n")

                elif module_type in ["schemas", "db-schema", "db_schemas"]:
                    tname = getattr(item, "table_name", "untitled_table")
                    lines.append(f"## Table: `{tname}`")
                    lines.append(f"**Status:** {getattr(item, 'status', 'draft').upper()}\n")
                    if getattr(item, "description", ""):
                        lines.append(f"{getattr(item, 'description')}\n")
                    cols = getattr(item, "columns", "[]")
                    try:
                        cols_list = json.loads(cols) if isinstance(cols, str) else cols
                        if cols_list and isinstance(cols_list, list):
                            lines.append("### Columns")
                            lines.append("| Column | Type | Nullable | Primary Key | Description |")
                            lines.append("| :--- | :--- | :--- | :--- | :--- |")
                            for col in cols_list:
                                lines.append(f"| `{col.get('name', col.get('column_name', ''))}` | {col.get('type', col.get('data_type', ''))} | {str(col.get('nullable', True))} | {str(col.get('primary_key', False))} | {col.get('description', '')} |")
                            lines.append("")
                    except Exception:
                        pass
                    diag = getattr(item, "mermaid_diagram", None)
                    if diag:
                        lines.append("### ER Diagram\n```mermaid\n" + diag.strip() + "\n```\n")
                    lines.append("---\n")

            buffer = io.BytesIO()
            buffer.write("\n".join(lines).encode("utf-8"))
            buffer.seek(0)
            return buffer, "text/markdown", "md"

        elif format_type == "docx":
            doc_obj = Document()
            doc_obj.add_heading(f"{project_name} - {module_type.replace('-', ' ').title()} Export", 0)
            doc_obj.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

            if not items:
                doc_obj.add_paragraph("No items found.")
            else:
                for item in items:
                    if module_type in ["designs", "system-design", "system_design"]:
                        doc_obj.add_heading(f"{getattr(item, 'title', 'Untitled Design')} ({getattr(item, 'architecture_type', 'System')})", level=1)
                        p = doc_obj.add_paragraph()
                        p.add_run(f"Status: {getattr(item, 'status', 'draft').upper()} | Version: {getattr(item, 'version', 1)}").bold = True
                        if getattr(item, "description", ""):
                            doc_obj.add_heading("Description", level=2)
                            doc_obj.add_paragraph(getattr(item, "description"))
                        tech = getattr(item, "tech_stack", "{}")
                        try:
                            tech_dict = json.loads(tech) if isinstance(tech, str) else tech
                            if tech_dict and isinstance(tech_dict, dict):
                                doc_obj.add_heading("Recommended Tech Stack", level=2)
                                for k, v in tech_dict.items():
                                    p = doc_obj.add_paragraph(style='List Bullet')
                                    p.add_run(f"{k.title()}: ").bold = True
                                    p.add_run(str(v))
                        except Exception:
                            pass
                        comps = getattr(item, "components", "[]")
                        try:
                            comps_list = json.loads(comps) if isinstance(comps, str) else comps
                            if comps_list and isinstance(comps_list, list):
                                doc_obj.add_heading("Architectural Components", level=2)
                                t = doc_obj.add_table(rows=1 + len(comps_list), cols=4)
                                t.style = 'Table Grid'
                                hdr = t.rows[0].cells
                                hdr[0].text = "Name"; hdr[1].text = "Type"; hdr[2].text = "Tech"; hdr[3].text = "Description"
                                for i, c in enumerate(comps_list):
                                    row = t.rows[i+1].cells
                                    row[0].text = str(c.get('name', ''))
                                    row[1].text = str(c.get('type', ''))
                                    row[2].text = str(c.get('tech', ''))
                                    row[3].text = str(c.get('description', ''))
                        except Exception:
                            pass
                        doc_obj.add_paragraph()

                    elif module_type in ["apis", "api-contracts", "api_contracts"]:
                        method = getattr(item, "method", "GET").upper()
                        path = getattr(item, "path", "/")
                        doc_obj.add_heading(f"[{method}] {path}", level=1)
                        p = doc_obj.add_paragraph()
                        p.add_run(f"Summary: {getattr(item, 'summary', '')} | Status: {getattr(item, 'status', 'draft').upper()}").bold = True
                        if getattr(item, "description", ""):
                            doc_obj.add_paragraph(getattr(item, "description"))
                        req = getattr(item, "request_body", None)
                        if req and req != "null":
                            doc_obj.add_heading("Request Body", level=2)
                            doc_obj.add_paragraph(str(req))
                        res = getattr(item, "response_body", None)
                        if res and res != "null":
                            doc_obj.add_heading("Response Body", level=2)
                            doc_obj.add_paragraph(str(res))
                        doc_obj.add_paragraph()

                    elif module_type in ["schemas", "db-schema", "db_schemas"]:
                        tname = getattr(item, "table_name", "untitled_table")
                        doc_obj.add_heading(f"Table: {tname}", level=1)
                        p = doc_obj.add_paragraph()
                        p.add_run(f"Status: {getattr(item, 'status', 'draft').upper()}").bold = True
                        if getattr(item, "description", ""):
                            doc_obj.add_paragraph(getattr(item, "description"))
                        cols = getattr(item, "columns", "[]")
                        try:
                            cols_list = json.loads(cols) if isinstance(cols, str) else cols
                            if cols_list and isinstance(cols_list, list):
                                doc_obj.add_heading("Columns", level=2)
                                t = doc_obj.add_table(rows=1 + len(cols_list), cols=5)
                                t.style = 'Table Grid'
                                hdr = t.rows[0].cells
                                hdr[0].text = "Column"; hdr[1].text = "Type"; hdr[2].text = "Nullable"; hdr[3].text = "PK"; hdr[4].text = "Description"
                                for i, c in enumerate(cols_list):
                                    row = t.rows[i+1].cells
                                    row[0].text = str(c.get('name', c.get('column_name', '')))
                                    row[1].text = str(c.get('type', c.get('data_type', '')))
                                    row[2].text = str(c.get('nullable', True))
                                    row[3].text = str(c.get('primary_key', False))
                                    row[4].text = str(c.get('description', ''))
                        except Exception:
                            pass
                        doc_obj.add_paragraph()

            buffer = io.BytesIO()
            doc_obj.save(buffer)
            buffer.seek(0)
            return buffer, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"

        else: # Default to PDF
            buffer = io.BytesIO()
            doc_pdf = SimpleDocTemplate(buffer, pagesize=landscape(A4), rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
            elements = []
            styles = getSampleStyleSheet()
            title_style = styles["Title"]
            normal_style = styles["Normal"]
            h1_style = styles["Heading1"]

            elements.append(Paragraph(f"{project_name} - {module_type.replace('-', ' ').title()} Export", title_style))
            elements.append(Paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", normal_style))
            elements.append(Spacer(1, 20))

            if not items:
                elements.append(Paragraph("No items found.", normal_style))
            else:
                for item in items:
                    if module_type in ["designs", "system-design", "system_design"]:
                        elements.append(Paragraph(f"<b>{getattr(item, 'title', 'Design')}</b> ({getattr(item, 'architecture_type', '')})", h1_style))
                        if getattr(item, 'description', ''):
                            elements.append(Paragraph(getattr(item, 'description'), normal_style))
                            elements.append(Spacer(1, 10))
                        comps = getattr(item, "components", "[]")
                        try:
                            comps_list = json.loads(comps) if isinstance(comps, str) else comps
                            if comps_list and isinstance(comps_list, list):
                                data = [["Name", "Type", "Tech", "Description"]]
                                for c in comps_list:
                                    data.append([
                                        Paragraph(str(c.get("name", "")), normal_style),
                                        Paragraph(str(c.get("type", "")), normal_style),
                                        Paragraph(str(c.get("tech", "")), normal_style),
                                        Paragraph(str(c.get("description", "")), normal_style),
                                    ])
                                table = Table(data, colWidths=[130, 100, 150, 400], repeatRows=1)
                                table.setStyle(TableStyle([
                                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#f1f5f9")),
                                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                    ('GRID', (0, 0), (-1, -1), 1, colors.HexColor("#e2e8f0")),
                                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                                ]))
                                elements.append(table)
                                elements.append(Spacer(1, 15))
                        except Exception:
                            pass

                    elif module_type in ["apis", "api-contracts", "api_contracts"]:
                        method = getattr(item, "method", "GET").upper()
                        path = getattr(item, "path", "/")
                        elements.append(Paragraph(f"<b>[{method}]</b> {path}", h1_style))
                        elements.append(Paragraph(f"<b>Summary:</b> {getattr(item, 'summary', '')}", normal_style))
                        if getattr(item, "description", ""):
                            elements.append(Paragraph(getattr(item, "description"), normal_style))
                        elements.append(Spacer(1, 15))

                    elif module_type in ["schemas", "db-schema", "db_schemas"]:
                        tname = getattr(item, "table_name", "untitled")
                        elements.append(Paragraph(f"<b>Table:</b> {tname}", h1_style))
                        if getattr(item, "description", ""):
                            elements.append(Paragraph(getattr(item, "description"), normal_style))
                            elements.append(Spacer(1, 10))
                        cols = getattr(item, "columns", "[]")
                        try:
                            cols_list = json.loads(cols) if isinstance(cols, str) else cols
                            if cols_list and isinstance(cols_list, list):
                                data = [["Column", "Type", "Null", "PK", "Description"]]
                                for col in cols_list:
                                    data.append([
                                        Paragraph(str(col.get("name", col.get("column_name", ""))), normal_style),
                                        Paragraph(str(col.get("type", col.get("data_type", ""))), normal_style),
                                        Paragraph(str(col.get("nullable", True)), normal_style),
                                        Paragraph(str(col.get("primary_key", False)), normal_style),
                                        Paragraph(str(col.get("description", "")), normal_style),
                                    ])
                                table = Table(data, colWidths=[150, 100, 50, 50, 432], repeatRows=1)
                                table.setStyle(TableStyle([
                                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#f1f5f9")),
                                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                    ('GRID', (0, 0), (-1, -1), 1, colors.HexColor("#e2e8f0")),
                                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                                ]))
                                elements.append(table)
                                elements.append(Spacer(1, 15))
                        except Exception:
                            pass

            doc_pdf.build(elements)
            buffer.seek(0)
            return buffer, "application/pdf", "pdf"

export_service = ExportService()
